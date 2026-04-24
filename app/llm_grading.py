from __future__ import annotations

import base64
import hashlib
import io
import json
import os
import subprocess
import tempfile
import mimetypes
import re
import threading
import time
from concurrent.futures import ThreadPoolExecutor, wait
import zipfile
from dataclasses import dataclass
from datetime import datetime, timedelta, timezone
from pathlib import Path, PurePosixPath
from typing import Any, Optional
from urllib.parse import urljoin, urlparse
from zoneinfo import ZoneInfo

import fitz
import httpx
from PIL import Image, ImageFile, UnidentifiedImageError
from docx import Document
import olefile
import openpyxl
import rarfile
import xlrd
from sqlalchemy import and_, func, text
from sqlalchemy.orm.attributes import flag_modified


def _unrar_tool_path() -> Optional[str]:
    import shutil

    return shutil.which("unrar") or shutil.which("unrar-free")


def _rar_read_member_bytes(archive_path: str, member_name: str) -> bytes:
    """Extract one member via unrar/unrar-free (RAR5 solid archives; 7z cannot read these)."""
    import shutil

    tool = _unrar_tool_path()
    if not tool:
        raise RuntimeError("未找到 unrar / unrar-free，无法解压 RAR。")
    abs_arc = os.path.abspath(archive_path)
    norm_member = (member_name or "").replace("\\", "/")
    tmp_dir = tempfile.mkdtemp(prefix="rar-one-")
    try:
        proc = subprocess.run(
            [tool, "x", "-o+", "-y", abs_arc, norm_member],
            cwd=tmp_dir,
            capture_output=True,
            timeout=120,
        )
        if proc.returncode != 0:
            err = (proc.stderr or proc.stdout or b"").decode("utf-8", errors="replace")[:500]
            raise RuntimeError(f"unrar 解压失败（{proc.returncode}）：{err}")
        out_path = os.path.join(tmp_dir, norm_member)
        if not os.path.isfile(out_path):
            base = os.path.basename(norm_member)
            alt = os.path.join(tmp_dir, base)
            out_path = alt if os.path.isfile(alt) else out_path
        if not os.path.isfile(out_path):
            raise RuntimeError("unrar 解压后未找到目标文件。")
        return Path(out_path).read_bytes()
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)
from sqlalchemy.exc import IntegrityError
from sqlalchemy.orm import Session, joinedload

from app.attachments import get_attachment_file_path
from app.config import settings
from app.llm_token_quota import (
    quota_calendar,
    resolve_effective_daily_student_tokens,
    resolve_max_parallel_grading_tasks,
)
from app.database import SessionLocal, engine
from app.llm_group_routing import GroupRoutingContext
from app.models import (
    CourseLLMConfig,
    CourseLLMConfigEndpoint,
    Homework,
    HomeworkAttempt,
    HomeworkGradingTask,
    HomeworkScoreCandidate,
    HomeworkSubmission,
    LLMEndpointPreset,
    LLMGroup,
    LLMQuotaReservation,
    LLMTokenUsageLog,
)

VISION_TEST_IMAGE_DATA_URL = (
    "data:image/png;base64,"
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8"
    "/w8AAusB9Y9nKXUAAAAASUVORK5CYII="
)
# Cap encoded image size for multi-modal test requests (avoids huge payloads to LLM APIs).
MAX_VISION_TEST_IMAGE_BYTES = 5 * 1024 * 1024


def build_png_data_url_from_image_bytes(data: bytes) -> str:
    """Load common formats (jpeg/png/webp/gif/bmp) and emit an OpenAI-compatible data:image/png;base64,... URL."""
    if not data or len(data) > MAX_VISION_TEST_IMAGE_BYTES:
        raise ValueError(f"Image must be non-empty and at most {MAX_VISION_TEST_IMAGE_BYTES} bytes.")
    prev_truncated = ImageFile.LOAD_TRUNCATED_IMAGES
    ImageFile.LOAD_TRUNCATED_IMAGES = True
    try:
        im = Image.open(io.BytesIO(data))
        im.load()
    except (UnidentifiedImageError, OSError) as exc:
        raise ValueError("无法将文件识别为支持的图片（请使用 JPEG/PNG/WebP 等）。") from exc
    finally:
        ImageFile.LOAD_TRUNCATED_IMAGES = prev_truncated
    if im.mode not in ("RGB", "RGBA"):
        im = im.convert("RGBA" if (getattr(im, "info", None) and im.info.get("transparency") is not None) else "RGB")
    out = io.BytesIO()
    im.save(out, format="PNG", optimize=True)
    raw = out.getvalue()
    if len(raw) > MAX_VISION_TEST_IMAGE_BYTES:
        raise ValueError("转码为 PNG 后仍过大，请使用更小的图片。")
    b64 = base64.b64encode(raw).decode("ascii")
    return f"data:image/png;base64,{b64}"


JSON_FENCE_PATTERN = re.compile(r"^\s*```(?:json)?\s*(.*?)\s*```\s*$", re.DOTALL | re.IGNORECASE)
MAX_ZIP_DEPTH = 4
MAX_ZIP_FILES = 100
MAX_ZIP_TOTAL_BYTES = 80 * 1024 * 1024
MAX_FILE_TEXT_CHARS = 12000
MAX_IPYNB_OUTPUT_CHARS = 6000
# Prior attempts included in the grading prompt (text-only summary); older rounds omitted for token cost.
ITERATION_CONTEXT_MAX_PRIOR_ATTEMPTS = 2
ITERATION_PRIOR_NOTE_CHAR_MAX = 900
ITERATION_PRIOR_COMMENT_CHAR_MAX = 500
SUPPORTED_TEXT_EXTENSIONS = {
    ".c",
    ".cc",
    ".cpp",
    ".csv",
    ".docx",
    ".go",
    ".h",
    ".hpp",
    ".html",
    ".java",
    ".js",
    ".json",
    ".jsx",
    ".md",
    ".py",
    ".rb",
    ".rs",
    ".sql",
    ".tex",
    ".ts",
    ".tsx",
    ".txt",
    ".vue",
    ".xml",
    ".yaml",
    ".yml",
}
SUPPORTED_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"}
ZIP_EXTENSIONS = {".zip"}
RAR_EXTENSIONS = {".rar"}
EXCEL_XLSX_EXTENSIONS = {".xlsx"}
EXCEL_XLS_EXTENSIONS = {".xls"}
LEGACY_WORD_DOC_EXTENSIONS = {".doc"}
IPYNB_EXTENSIONS = {".ipynb"}
PDF_EXTENSIONS = {".pdf"}
RETRYABLE_STATUS_CODES = {408, 409, 425, 429, 500, 502, 503, 504}
NON_RETRYABLE_STATUS_CODES = {401, 403}

# After the first failed grading task for an attempt, enqueue at most this many extra tries.
_MAX_AUTO_RETRY_TASKS_PER_ATTEMPT = 2
_LLM_CALL_LOG_MAX_EVENTS = 60

_AUTO_RETRY_ELIGIBLE_ERROR_CODES = frozenset(
    {
        "llm_call_failed",
        "unexpected_error",
    }
)


@dataclass
class MaterialBlock:
    priority: int
    path: str
    block_type: str
    text: Optional[str] = None
    image_data_url: Optional[str] = None
    estimated_tokens: int = 0
    logical_path: Optional[str] = None
    mime_hint: Optional[str] = None
    origin: Optional[str] = None
    truncated: bool = False


class RetryableLLMError(Exception):
    pass


class NonRetryableLLMError(Exception):
    pass


class _WorkerManager:
    def __init__(self) -> None:
        self._thread: Optional[threading.Thread] = None
        self._stop_event = threading.Event()
        self._lock = threading.Lock()
        self._executor: Optional[ThreadPoolExecutor] = None
        self._executor_workers: int = 0

    def start(self) -> None:
        if not settings.ENABLE_LLM_GRADING_WORKER:
            return
        with self._lock:
            if self._thread and self._thread.is_alive():
                return
            self._stop_event.clear()
            self._thread = threading.Thread(target=self._run, name="llm-grading-worker", daemon=True)
            self._thread.start()

    def stop(self) -> None:
        self._stop_event.set()
        with self._lock:
            if self._executor:
                self._executor.shutdown(wait=False)
                self._executor = None
                self._executor_workers = 0

    def _ensure_executor(self, workers: int) -> ThreadPoolExecutor:
        w = max(1, int(workers))
        if self._executor is not None and self._executor_workers == w:
            return self._executor
        if self._executor is not None:
            self._executor.shutdown(wait=False, cancel_futures=True)
        self._executor = ThreadPoolExecutor(max_workers=w, thread_name_prefix="llm-grade")
        self._executor_workers = w
        return self._executor

    def _run(self) -> None:
        while not self._stop_event.is_set():
            try:
                db = SessionLocal()
                try:
                    cap = resolve_max_parallel_grading_tasks(db)
                finally:
                    db.close()
                claimed = claim_grading_tasks_batch(cap)
                if not claimed:
                    self._stop_event.wait(settings.LLM_GRADING_WORKER_POLL_SECONDS)
                    continue
                ex = self._ensure_executor(cap)
                futs = [ex.submit(process_grading_task, tid) for tid in claimed]
                wait(futs)
                for fut in futs:
                    try:
                        fut.result()
                    except Exception as exc:  # pragma: no cover
                        print(f"LLM grading task error: {exc}")
            except Exception as exc:  # pragma: no cover - defensive background logging
                print(f"LLM grading worker loop error: {exc}")
                self._stop_event.wait(settings.LLM_GRADING_WORKER_POLL_SECONDS)


worker_manager = _WorkerManager()

# Serialize grading for a single task id in-process (duplicate worker wakeups / tests).
_task_grading_locks: dict[int, threading.Lock] = {}
_task_grading_locks_mutex = threading.Lock()


def _grading_lock_for_task(task_id: int) -> threading.Lock:
    with _task_grading_locks_mutex:
        if task_id not in _task_grading_locks:
            _task_grading_locks[task_id] = threading.Lock()
        return _task_grading_locks[task_id]


def start_grading_worker() -> None:
    worker_manager.start()


def _now_utc() -> datetime:
    return datetime.now(timezone.utc)


def normalize_score_for_homework(homework: Homework, score: float | int) -> float:
    value = max(0.0, min(float(score), float(homework.max_score or 100)))
    if (homework.grade_precision or "integer") == "decimal_1":
        return round(value, 1)
    return float(round(value))


def _teacher_candidate_sort_key(candidate: HomeworkScoreCandidate) -> tuple[float, datetime]:
    """Higher score first, then newer; used only among teacher rows."""
    ts = candidate.updated_at or candidate.created_at or datetime.min.replace(tzinfo=timezone.utc)
    return (float(candidate.score or 0), ts)


def _auto_candidate_sort_key(candidate: HomeworkScoreCandidate) -> tuple[float, datetime]:
    """Among auto rows: higher score, then newer."""
    ts = candidate.updated_at or candidate.created_at or datetime.min.replace(tzinfo=timezone.utc)
    return (float(candidate.score or 0), ts)


def get_best_score_candidate(
    db: Session,
    homework_id: int,
    student_id: int,
    *,
    latest_attempt_id: Optional[int] = None,
) -> Optional[HomeworkScoreCandidate]:
    """
    Best visible score for the submission summary: only the **latest attempt**'s
    candidates apply (if latest_attempt_id is set). On that attempt, any **teacher**
    score takes precedence over automatic scores so LLM output cannot override a
    teacher's grade in the UI.
    """
    query = (
        db.query(HomeworkScoreCandidate)
        .filter(
            HomeworkScoreCandidate.homework_id == homework_id,
            HomeworkScoreCandidate.student_id == student_id,
        )
    )
    if latest_attempt_id is not None:
        query = query.filter(HomeworkScoreCandidate.attempt_id == latest_attempt_id)
    candidates = query.all()
    valid_candidates = [
        c
        for c in candidates
        if c.score is not None and getattr(c.attempt, "counts_toward_final_score", True)
    ]
    if not valid_candidates:
        return None
    teacher_rows = [c for c in valid_candidates if c.source == "teacher"]
    if teacher_rows:
        return max(teacher_rows, key=_teacher_candidate_sort_key)
    auto_rows = [c for c in valid_candidates if c.source == "auto"]
    if auto_rows:
        return max(auto_rows, key=_auto_candidate_sort_key)
    return max(valid_candidates, key=_auto_candidate_sort_key)


def refresh_submission_summary(db: Session, summary: HomeworkSubmission) -> HomeworkSubmission:
    best_candidate = get_best_score_candidate(
        db, summary.homework_id, summary.student_id, latest_attempt_id=summary.latest_attempt_id
    )
    summary.review_score = best_candidate.score if best_candidate else None
    summary.review_comment = (best_candidate.comment or None) if best_candidate else None

    if summary.latest_attempt_id:
        latest_attempt = (
            db.query(HomeworkAttempt)
            .filter(HomeworkAttempt.id == summary.latest_attempt_id)
            .first()
        )
        if latest_attempt:
            summary.content = latest_attempt.content
            summary.attachment_name = latest_attempt.attachment_name
            summary.attachment_url = latest_attempt.attachment_url
            summary.submitted_at = latest_attempt.submitted_at
            latest_task = (
                db.query(HomeworkGradingTask)
                .filter(HomeworkGradingTask.attempt_id == latest_attempt.id)
                .order_by(HomeworkGradingTask.created_at.desc(), HomeworkGradingTask.id.desc())
                .first()
            )
            summary.latest_task_status = latest_task.status if latest_task else None
            err_task = latest_task
            if latest_task and latest_task.status in ("queued", "processing"):
                prev_failed = (
                    db.query(HomeworkGradingTask)
                    .filter(
                        HomeworkGradingTask.attempt_id == latest_attempt.id,
                        HomeworkGradingTask.status == "failed",
                        HomeworkGradingTask.id < latest_task.id,
                    )
                    .order_by(HomeworkGradingTask.id.desc())
                    .first()
                )
                if prev_failed:
                    err_task = prev_failed
            summary.latest_task_error = err_task.error_message if err_task else None
    return summary


def ensure_course_llm_config(db: Session, subject_id: int, user_id: Optional[int] = None) -> CourseLLMConfig:
    config = db.query(CourseLLMConfig).filter(CourseLLMConfig.subject_id == subject_id).first()
    if config:
        return config
    config = CourseLLMConfig(subject_id=subject_id, created_by=user_id, updated_by=user_id)
    db.add(config)
    db.flush()
    return config


def build_task_summary(task: HomeworkGradingTask) -> str:
    status_map = {
        "queued": "排队中",
        "processing": "处理中",
        "success": "成功",
        "failed": "失败",
    }
    status_label = status_map.get(task.status, task.status)
    if task.error_message:
        return f"{status_label}: {task.error_message}"
    return status_label


def _reclaim_stale_processing_tasks(db: Session) -> int:
    stale_before = _now_utc() - timedelta(seconds=max(60, int(settings.LLM_GRADING_TASK_STALE_SECONDS or 600)))
    stale_tasks = (
        db.query(HomeworkGradingTask)
        .filter(
            HomeworkGradingTask.status == "processing",
            HomeworkGradingTask.updated_at.isnot(None),
            HomeworkGradingTask.updated_at < stale_before,
        )
        .all()
    )
    reclaimed = 0
    for task in stale_tasks:
        task.status = "queued"
        task.error_code = None
        task.error_message = None
        task.queue_reason = "reclaimed_stale_processing"
        task.task_summary = "已回收超时任务，等待重试"
        task.started_at = None
        task.finished_at = None
        task.updated_at = _now_utc()
        reclaimed += 1
    if reclaimed:
        db.commit()
    return reclaimed


def queue_grading_task(
    db: Session,
    attempt: HomeworkAttempt,
    queue_reason: str = "new_submission",
) -> HomeworkGradingTask:
    existing_task = (
        db.query(HomeworkGradingTask)
        .filter(
            HomeworkGradingTask.attempt_id == attempt.id,
            HomeworkGradingTask.status.in_(("queued", "processing")),
        )
        .order_by(HomeworkGradingTask.created_at.desc(), HomeworkGradingTask.id.desc())
        .first()
    )
    if existing_task:
        summary = (
            db.query(HomeworkSubmission)
            .filter(
                HomeworkSubmission.homework_id == attempt.homework_id,
                HomeworkSubmission.student_id == attempt.student_id,
            )
            .first()
        )
        if summary:
            summary.latest_task_status = existing_task.status
            summary.latest_task_error = existing_task.error_message
            refresh_submission_summary(db, summary)
        return existing_task

    task = HomeworkGradingTask(
        attempt_id=attempt.id,
        homework_id=attempt.homework_id,
        student_id=attempt.student_id,
        subject_id=attempt.subject_id,
        status="queued",
        queue_reason=queue_reason,
    )
    db.add(task)
    db.flush()
    summary = (
        db.query(HomeworkSubmission)
        .filter(
            HomeworkSubmission.homework_id == attempt.homework_id,
            HomeworkSubmission.student_id == attempt.student_id,
        )
        .first()
    )
    if summary:
        summary.latest_task_status = task.status
        summary.latest_task_error = None
        refresh_submission_summary(db, summary)
    return task


def _append_llm_call_log(task: HomeworkGradingTask, event: dict[str, Any]) -> None:
    if not isinstance(task.artifact_manifest, dict):
        task.artifact_manifest = {}
    log = task.artifact_manifest.get("llm_call_log")
    if not isinstance(log, list):
        log = []
    event = {**event, "ts": _now_utc().isoformat()}
    log.append(event)
    if len(log) > _LLM_CALL_LOG_MAX_EVENTS:
        log = log[-_LLM_CALL_LOG_MAX_EVENTS :]
    task.artifact_manifest["llm_call_log"] = log
    flag_modified(task, "artifact_manifest")


def _flag_artifact_manifest_modified(task: HomeworkGradingTask) -> None:
    """JSON columns need explicit dirty flag when mutating nested dicts in place."""
    flag_modified(task, "artifact_manifest")


def _material_needs_vision(material: dict[str, Any]) -> bool:
    for block in material.get("student_blocks") or []:
        if getattr(block, "block_type", None) == "image":
            return True
    return False


def _preset_text_ready(preset: Optional[LLMEndpointPreset]) -> bool:
    if not preset or not preset.is_active:
        return False
    if preset.validation_status != "validated":
        return False
    ts = getattr(preset, "text_validation_status", None)
    if ts == "failed":
        return False
    # Legacy rows may only set validation_status; treat unknown as OK if overall validated.
    return ts in (None, "passed", "skipped")


def _preset_eligible_for_grading(preset: Optional[LLMEndpointPreset], *, need_vision: bool) -> tuple[bool, str]:
    if not preset:
        return False, "端点预设不存在。"
    if not preset.is_active:
        return False, f"端点「{preset.name}」已停用。"
    if preset.validation_status != "validated":
        return False, f"端点「{preset.name}」未通过整体校验（状态：{preset.validation_status}）。"
    if not _preset_text_ready(preset):
        msg = getattr(preset, "text_validation_message", None) or "未通过纯文本连通性测试"
        return False, f"端点「{preset.name}」{msg}。"
    if need_vision:
        if not preset.supports_vision:
            return False, f"端点「{preset.name}」未声明支持视觉，无法处理含图片/PDF 的提交。"
        vs = getattr(preset, "vision_validation_status", None)
        if vs == "failed":
            vm = getattr(preset, "vision_validation_message", None) or "视觉连通性未通过"
            return False, f"端点「{preset.name}」{vm}。"
        if vs not in (None, "passed", "skipped"):
            vm = getattr(preset, "vision_validation_message", None) or "视觉连通性未通过"
            return False, f"端点「{preset.name}」{vm}。"
    return True, ""


def _homework_routing_warnings(db: Session, homework: Homework, config: CourseLLMConfig) -> list[str]:
    """Return user-facing warnings when homework.llm_routing_spec may diverge from course defaults."""
    spec = homework.llm_routing_spec
    if not spec or not isinstance(spec, dict):
        return []
    mode = spec.get("mode")
    warnings: list[str] = []
    if mode == "latest_passing_validated":
        preset = (
            db.query(LLMEndpointPreset)
            .filter(
                LLMEndpointPreset.is_active.is_(True),
                LLMEndpointPreset.validation_status == "validated",
                LLMEndpointPreset.text_validation_status == "passed",
            )
            .order_by(LLMEndpointPreset.validated_at.desc().nullslast(), LLMEndpointPreset.id.desc())
            .first()
        )
        if not preset:
            warnings.append("作业要求使用「最新纯文本测试通过」的端点，但系统中没有符合条件的预设，已按课程设置路由。")
            return warnings
        bound_ids = {m.preset_id for g in (config.groups or []) for m in (g.members or [])}
        bound_ids |= {e.preset_id for e in (config.endpoints or [])}
        if preset.id not in bound_ids:
            warnings.append(
                f"作业要求优先使用「{preset.name}」，但该预设未加入本课程的 LLM 配置，调用仍可能失败。"
            )
    if mode == "limit_to_preset_ids":
        raw_ids = spec.get("preset_ids")
        if isinstance(raw_ids, list) and raw_ids:
            id_set: set[int] = set()
            for x in raw_ids:
                try:
                    id_set.add(int(x))
                except (TypeError, ValueError):
                    continue
            if id_set:
                any_hit = any(
                    m.preset_id in id_set for g in (config.groups or []) for m in (g.members or [])
                ) or any(e.preset_id in id_set for e in (config.endpoints or []))
                if not any_hit:
                    warnings.append("作业限制了端点预设，但与课程设置无交集，已回退为课程完整路由。")
    return warnings


def _filter_course_links_for_homework(
    homework: Homework,
    group_rows: list[LLMGroup],
    flat_endpoints: list[CourseLLMConfigEndpoint],
) -> tuple[list[Any], list[CourseLLMConfigEndpoint], list[str]]:
    """Apply homework.llm_routing_spec (preset subset) without mutating ORM collections."""
    from types import SimpleNamespace

    spec = homework.llm_routing_spec
    notes: list[str] = []
    if not spec or not isinstance(spec, dict):
        return group_rows, flat_endpoints, notes

    mode = spec.get("mode")
    if mode != "limit_to_preset_ids":
        return group_rows, flat_endpoints, notes

    raw_ids = spec.get("preset_ids")
    if not isinstance(raw_ids, list) or not raw_ids:
        return group_rows, flat_endpoints, notes
    id_set: set[int] = set()
    for x in raw_ids:
        try:
            id_set.add(int(x))
        except (TypeError, ValueError):
            continue
    if not id_set:
        return group_rows, flat_endpoints, notes

    ephemeral_groups: list[Any] = []
    for g in sorted(group_rows, key=lambda x: (x.priority, x.id)):
        members = [m for m in (g.members or []) if m.preset_id in id_set]
        if members:
            ephemeral_groups.append(
                SimpleNamespace(id=g.id, priority=g.priority, name=getattr(g, "name", None), members=members)
            )
    new_flat = [e for e in flat_endpoints if e.preset_id in id_set]
    if ephemeral_groups:
        notes.append("routing_mode:limit_to_preset_ids(groups)")
        return ephemeral_groups, [], notes
    if new_flat:
        notes.append("routing_mode:limit_to_preset_ids(flat)")
        return [], new_flat, notes
    notes.append("routing_mode:limit_to_preset_ids_miss")
    return group_rows, flat_endpoints, notes


def validate_text_connectivity(
    base_url: str,
    api_key: str,
    model_name: str,
    connect_timeout_seconds: int,
    read_timeout_seconds: int,
) -> tuple[bool, str]:
    """OpenAI-style chat: text-only smoke test before multimodal check."""
    timeout = httpx.Timeout(connect=connect_timeout_seconds, read=read_timeout_seconds, write=read_timeout_seconds, pool=connect_timeout_seconds)
    messages = [
        {
            "role": "user",
            "content": "Please reply with the single word: OK",
        }
    ]
    payload = {
        "model": model_name,
        "messages": messages,
        "temperature": 0,
        "max_tokens": 8,
    }
    endpoint_url = _build_chat_completion_url(base_url)
    try:
        with httpx.Client(timeout=timeout) as client:
            response = client.post(
                endpoint_url,
                json=payload,
                headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            )
    except httpx.HTTPError as exc:
        return False, f"纯文本连通性校验失败：{exc}"

    if response.status_code >= 400:
        return False, f"纯文本连通性校验失败：HTTP {response.status_code} {response.text[:300]}"

    try:
        data = response.json()
    except ValueError:
        return False, "纯文本连通性校验失败：返回内容不是 JSON。"

    content = _extract_message_content(data)
    if not content.strip():
        return False, "纯文本连通性校验失败：模型未返回可读文本。"

    return True, "纯文本请求校验通过。"


def validate_vision_connectivity(
    base_url: str,
    api_key: str,
    model_name: str,
    connect_timeout_seconds: int,
    read_timeout_seconds: int,
    image_data_url: Optional[str] = None,
) -> tuple[bool, str]:
    if image_data_url and not (image_data_url.startswith("data:image/") and "base64," in image_data_url):
        return False, "视觉能力校验失败：图片数据格式无效（需为 data:image/...;base64,...）。"
    data_url = image_data_url or VISION_TEST_IMAGE_DATA_URL
    timeout = httpx.Timeout(connect=connect_timeout_seconds, read=read_timeout_seconds, write=read_timeout_seconds, pool=connect_timeout_seconds)
    messages = [
        {
            "role": "user",
            "content": [
                {"type": "text", "text": "Please reply with OK."},
                {"type": "image_url", "image_url": {"url": data_url}},
            ],
        }
    ]
    payload = {
        "model": model_name,
        "messages": messages,
        "temperature": 0,
        "max_tokens": 5,
    }

    endpoint_url = _build_chat_completion_url(base_url)
    try:
        with httpx.Client(timeout=timeout) as client:
            response = client.post(
                endpoint_url,
                json=payload,
                headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            )
    except httpx.HTTPError as exc:
        return False, f"视觉能力校验失败：{exc}"

    if response.status_code >= 400:
        return False, f"视觉能力校验失败：HTTP {response.status_code} {response.text[:300]}"

    try:
        data = response.json()
    except ValueError:
        return False, "视觉能力校验失败：返回内容不是 JSON。"

    content = _extract_message_content(data)
    if not content.strip():
        return False, "视觉能力校验失败：模型未返回可读文本。"

    return True, "多模态（图像）输入校验通过。"


def validate_endpoint_connectivity(
    base_url: str,
    api_key: str,
    model_name: str,
    connect_timeout_seconds: int,
    read_timeout_seconds: int,
) -> tuple[bool, str]:
    ok, msg = validate_text_connectivity(
        base_url, api_key, model_name, connect_timeout_seconds, read_timeout_seconds
    )
    if not ok:
        return False, msg
    ok2, msg2 = validate_vision_connectivity(
        base_url, api_key, model_name, connect_timeout_seconds, read_timeout_seconds
    )
    if not ok2:
        return False, msg2
    return True, "端点连通性校验通过：已验证纯文本与多模态（图像）输入。"


def estimate_task_tokens(
    config: CourseLLMConfig,
    text_length: int,
    image_count: int,
) -> int:
    chars_per_token = config.estimated_chars_per_token or 4.0
    text_tokens = int(text_length / chars_per_token) + 200
    image_tokens = int(image_count * (config.estimated_image_tokens or 850))
    return text_tokens + image_tokens + int(config.max_output_tokens or 0)


# Stable section markers for the scoring prompt (model + debugging)
SECTION_ASSIGNMENT = "[SECTION:INSTRUCTOR_ASSIGNMENT]"
SECTION_STUDENT_BODY = "[SECTION:STUDENT_TEXT_RESPONSE]"
SECTION_ATTACHMENT = "[SECTION:ATTACHMENT_CONTENT]"
SECTION_IMAGES = "[SECTION:STUDENT_IMAGES]"
SECTION_NOTES = "[SECTION:PIPELINE_NOTES]"


def _attachment_block_meta_text(block: "MaterialBlock") -> str:
    """Short human-readable line for prompts (no raw base64)."""
    if block.block_type != "text" or not (block.origin or "").startswith("attachment"):
        return ""
    name = (block.logical_path or block.path or "attachment").strip()
    mime = (block.mime_hint or "").strip()
    origin = (block.origin or "").strip()
    flags: list[str] = []
    if block.truncated:
        flags.append("truncated")
    flag_s = f" flags={','.join(flags)}" if flags else ""
    mime_s = f" mime={mime}" if mime else ""
    return f"[ATTACHMENT_META path={name} origin={origin}{mime_s}{flag_s}]\n"


def _data_url_payload_chars(data_url: Optional[str]) -> int:
    if not data_url:
        return 0
    if ";base64," in data_url:
        return max(0, len(data_url) - data_url.find(";base64,") - len(";base64,"))
    return len(data_url)


def estimate_request_tokens_from_material(
    config: CourseLLMConfig,
    material: dict[str, Any],
    *,
    assignment_text: str,
    teacher_prompt: str,
    student_intro: str,
) -> int:
    """
    Approximate request-side tokens using text char counts (incl. base64 payload length
    for image data URLs) plus configured image/output heuristics.
    """
    chars_per_token = float(config.estimated_chars_per_token or 4.0)
    per_image = int(config.estimated_image_tokens or 850)

    section_overhead = len(SECTION_ASSIGNMENT) + len(SECTION_STUDENT_BODY)
    section_overhead += len(SECTION_ATTACHMENT) + len(SECTION_IMAGES) + len(SECTION_NOTES)

    text_chars = (
        len(assignment_text or "")
        + len(teacher_prompt or "")
        + len(student_intro or "")
        + section_overhead
    )
    text_chars += len(material.get("notes_text") or "")

    for block in material.get("student_blocks") or []:
        if block.block_type == "text":
            text_chars += len(block.text or "")
            meta = _attachment_block_meta_text(block)
            if meta:
                text_chars += len(meta)
        elif block.block_type == "image":
            text_chars += _data_url_payload_chars(block.image_data_url)

    text_tokens = int(text_chars / chars_per_token) + 200
    image_count = len([b for b in (material.get("student_blocks") or []) if b.block_type == "image"])
    image_tokens = int(image_count * per_image)
    return text_tokens + image_tokens + int(config.max_output_tokens or 0)


def get_student_quota_usage_snapshot(db: Session, *, student_id: int) -> dict[str, Any]:
    """Per-user usage for quota day (student-facing UI; excludes API keys)."""
    usage_date, timezone_name = quota_calendar(db)
    lim_stu = resolve_effective_daily_student_tokens(db, student_id)
    snap: dict[str, Any] = {
        "usage_date": usage_date,
        "quota_timezone": timezone_name,
        "daily_student_token_limit": lim_stu,
        "student_used_tokens_today": None,
        "student_remaining_tokens_today": None,
    }
    used_stu = _get_used_tokens_for_scope(
        db,
        usage_date=usage_date,
        timezone_name=timezone_name,
        student_id=student_id,
    )
    used_stu += _sum_reserved_tokens(
        db,
        usage_date=usage_date,
        timezone_name=timezone_name,
        student_id=student_id,
    )
    snap["student_used_tokens_today"] = used_stu
    snap["student_remaining_tokens_today"] = max(0, lim_stu - used_stu)
    return snap


def _quota_delta_violations(
    db: Session,
    config: CourseLLMConfig,
    *,
    usage_date: str,
    timezone_name: str,
    student_id: int,
    subject_id: Optional[int],
    delta_tokens: int,
) -> list[str]:
    violations: list[str] = []
    lim_stu = resolve_effective_daily_student_tokens(db, student_id)
    used_by_student = _get_used_tokens_for_scope(
        db,
        usage_date=usage_date,
        timezone_name=timezone_name,
        student_id=student_id,
    )
    if used_by_student + delta_tokens > lim_stu:
        violations.append("student")
    return violations


def _get_usage_date(timezone_name: str) -> str:
    try:
        tz = ZoneInfo(timezone_name or "UTC")
    except Exception:
        tz = ZoneInfo("UTC")
    return datetime.now(tz).date().isoformat()


def _get_used_tokens_for_scope(
    db: Session,
    *,
    usage_date: str,
    timezone_name: str,
    student_id: Optional[int] = None,
    subject_id: Optional[int] = None,
) -> int:
    query = db.query(LLMTokenUsageLog).filter(
        LLMTokenUsageLog.usage_date == usage_date,
        LLMTokenUsageLog.timezone == timezone_name,
    )
    if student_id is not None:
        query = query.filter(LLMTokenUsageLog.student_id == student_id)
    if subject_id is not None:
        query = query.filter(LLMTokenUsageLog.subject_id == subject_id)
    total = 0
    for item in query.all():
        if item.total_tokens is not None:
            total += int(item.total_tokens)
        else:
            total += int(item.input_tokens or 0) + int(item.output_tokens or 0)
    return total


def _sum_reserved_tokens(
    db: Session,
    *,
    usage_date: str,
    timezone_name: str,
    student_id: Optional[int] = None,
    subject_id: Optional[int] = None,
) -> int:
    q = db.query(func.coalesce(func.sum(LLMQuotaReservation.reserved_tokens), 0)).filter(
        LLMQuotaReservation.usage_date == usage_date,
        LLMQuotaReservation.timezone == timezone_name,
    )
    if student_id is not None:
        q = q.filter(LLMQuotaReservation.student_id == student_id)
    if subject_id is not None:
        q = q.filter(LLMQuotaReservation.subject_id == subject_id)
    val = q.scalar()
    return int(val or 0)


def _hash_to_pg_advisory_key(label: str) -> int:
    digest = hashlib.sha256(label.encode("utf-8")).digest()[:8]
    return int.from_bytes(digest, "big", signed=False) % (2**62)


def _pg_quota_advisory_keys(
    *,
    student_id: int,
    usage_date: str,
    timezone_name: str,
    effective_student_daily_cap: int,
) -> list[int]:
    keys: list[int] = []
    if effective_student_daily_cap and effective_student_daily_cap > 0:
        keys.append(
            _hash_to_pg_advisory_key(f"llm_quota|student|{student_id}|{usage_date}|{timezone_name}")
        )
    return sorted(set(keys))


_quota_serialization_lock = threading.Lock()


def release_quota_reservation(db: Session, task_id: int) -> None:
    db.query(LLMQuotaReservation).filter(LLMQuotaReservation.task_id == task_id).delete(synchronize_session=False)


def _quota_precheck_in_session(
    db: Session,
    config: CourseLLMConfig,
    *,
    student_id: int,
    subject_id: Optional[int],
    estimated_tokens: int,
) -> tuple[bool, Optional[str]]:
    """Committed usage plus in-flight reservations vs daily caps."""
    usage_date, timezone_name = quota_calendar(db)
    lim_stu = resolve_effective_daily_student_tokens(db, student_id)
    used_by_student = _get_used_tokens_for_scope(
        db,
        usage_date=usage_date,
        timezone_name=timezone_name,
        student_id=student_id,
    )
    used_by_student += _sum_reserved_tokens(
        db,
        usage_date=usage_date,
        timezone_name=timezone_name,
        student_id=student_id,
    )
    if used_by_student + estimated_tokens > lim_stu:
        return False, "quota_exceeded_student"
    return True, None


def precheck_quota(
    db: Session,
    config: CourseLLMConfig,
    *,
    student_id: int,
    subject_id: Optional[int],
    estimated_tokens: int,
) -> tuple[bool, Optional[str]]:
    """Best-effort check including in-flight reservations (see reserve_quota_tokens for DB-locked reserve)."""
    return _quota_precheck_in_session(db, config, student_id=student_id, subject_id=subject_id, estimated_tokens=estimated_tokens)


def reserve_quota_tokens(
    db: Session,
    task: HomeworkGradingTask,
    config: CourseLLMConfig,
    estimated_tokens: int,
) -> tuple[bool, Optional[str]]:
    """
    Reserve estimated tokens against daily caps so concurrent workers share one budget.
    PostgreSQL: short transaction + pg_advisory_xact_lock per scope.
    SQLite / others: process-wide lock + same-session row (tests).
    """
    usage_date, timezone_name = quota_calendar(db)

    def _try_insert_reservation(sess: Session) -> tuple[bool, Optional[str]]:
        release_quota_reservation(sess, task.id)
        ok, err = _quota_precheck_in_session(
            sess,
            config,
            student_id=task.student_id,
            subject_id=task.subject_id,
            estimated_tokens=estimated_tokens,
        )
        if not ok:
            return ok, err
        sess.add(
            LLMQuotaReservation(
                task_id=task.id,
                student_id=task.student_id,
                subject_id=task.subject_id,
                usage_date=usage_date,
                timezone=timezone_name,
                reserved_tokens=int(estimated_tokens),
            )
        )
        sess.flush()
        return True, None

    if engine.dialect.name == "postgresql":
        keys = _pg_quota_advisory_keys(
            student_id=task.student_id,
            usage_date=usage_date,
            timezone_name=timezone_name,
            effective_student_daily_cap=resolve_effective_daily_student_tokens(db, task.student_id),
        )
        with engine.begin() as conn:
            for k in keys:
                conn.execute(text("SELECT pg_advisory_xact_lock(CAST(:k AS BIGINT))"), {"k": k})
            inner = Session(bind=conn, autoflush=False, autocommit=False)
            try:
                ok, err = _try_insert_reservation(inner)
                if not ok:
                    return ok, err
            except IntegrityError:
                return True, None
            finally:
                inner.close()
        return True, None

    with _quota_serialization_lock:
        return _try_insert_reservation(db)


def record_usage_if_needed(
    db: Session,
    task: HomeworkGradingTask,
    config: CourseLLMConfig,
    usage: dict[str, Any],
) -> None:
    existing = db.query(LLMTokenUsageLog).filter(LLMTokenUsageLog.task_id == task.id).first()
    if existing:
        return
    release_quota_reservation(db, task.id)
    usage_date, timezone_name = quota_calendar(db)
    prompt_tokens = usage.get("prompt_tokens")
    completion_tokens = usage.get("completion_tokens")
    total_tokens = usage.get("total_tokens")
    if total_tokens is None:
        total_tokens = int(prompt_tokens or 0) + int(completion_tokens or 0)
    billing_note: Optional[str] = None
    violations = _quota_delta_violations(
        db,
        config,
        usage_date=usage_date,
        timezone_name=timezone_name,
        student_id=task.student_id,
        subject_id=task.subject_id,
        delta_tokens=int(total_tokens or 0),
    )
    if violations:
        billing_note = "over_daily_limit:" + ",".join(violations)

    db.add(
        LLMTokenUsageLog(
            task_id=task.id,
            subject_id=task.subject_id,
            student_id=task.student_id,
            usage_date=usage_date,
            timezone=timezone_name,
            input_tokens=prompt_tokens,
            output_tokens=completion_tokens,
            total_tokens=total_tokens,
            billing_note=billing_note,
        )
    )
    task.billed_input_tokens = prompt_tokens
    task.billed_output_tokens = completion_tokens
    task.billed_total_tokens = total_tokens


def claim_grading_tasks_batch(max_tasks: int) -> list[int]:
    """
    Atomically move up to max_tasks rows from queued -> processing (oldest first).
    Returns list of task ids claimed in this transaction (empty if none).
    """
    if max_tasks < 1:
        return []
    db = SessionLocal()
    try:
        _reclaim_stale_processing_tasks(db)
        candidates = (
            db.query(HomeworkGradingTask)
            .filter(HomeworkGradingTask.status == "queued")
            .order_by(HomeworkGradingTask.created_at.asc(), HomeworkGradingTask.id.asc())
            .limit(max_tasks)
            .all()
        )
        if not candidates:
            return []
        now = _now_utc()
        claimed: list[int] = []
        for task in candidates:
            n = (
                db.query(HomeworkGradingTask)
                .filter(
                    HomeworkGradingTask.id == task.id,
                    HomeworkGradingTask.status == "queued",
                )
                .update(
                    {
                        HomeworkGradingTask.status: "processing",
                        HomeworkGradingTask.started_at: now,
                        HomeworkGradingTask.updated_at: now,
                        HomeworkGradingTask.task_summary: "处理中",
                    },
                    synchronize_session=False,
                )
            )
            if n:
                claimed.append(int(task.id))
        if not claimed:
            db.rollback()
            return []
        db.commit()
        return claimed
    except Exception:
        db.rollback()
        raise
    finally:
        db.close()


def process_next_grading_task() -> bool:
    claimed = claim_grading_tasks_batch(1)
    if not claimed:
        return False
    process_grading_task(claimed[0])
    return True


def process_grading_task(task_id: int) -> None:
    db = SessionLocal()
    try:
        _reclaim_stale_processing_tasks(db)
    finally:
        db.close()
    with _grading_lock_for_task(task_id):
        _process_grading_task_unlocked(task_id)


def _process_grading_task_unlocked(task_id: int) -> None:
    db = SessionLocal()
    try:
        task = db.query(HomeworkGradingTask).filter(HomeworkGradingTask.id == task_id).first()
        if not task:
            return
        if task.status in ("success", "failed"):
            return
        if task.status == "queued":
            # Claim the task (tests call process_grading_task directly; worker uses process_next which pre-sets processing).
            now = _now_utc()
            n = (
                db.query(HomeworkGradingTask)
                .filter(HomeworkGradingTask.id == task_id, HomeworkGradingTask.status == "queued")
                .update(
                    {
                        HomeworkGradingTask.status: "processing",
                        HomeworkGradingTask.started_at: now,
                        HomeworkGradingTask.updated_at: now,
                        HomeworkGradingTask.task_summary: "处理中",
                    },
                    synchronize_session=False,
                )
            )
            if not n:
                return
            db.commit()
            task = db.query(HomeworkGradingTask).filter(HomeworkGradingTask.id == task_id).first()
        elif task.status != "processing":
            return
        try:
            _run_grading_after_claim(db, task_id, task)
        except Exception as exc:
            db.rollback()
            task2 = db.query(HomeworkGradingTask).filter(HomeworkGradingTask.id == task_id).first()
            if task2:
                _mark_task_failed(db, task2, "unexpected_error", f"评分任务异常：{exc}")
    finally:
        db.close()


def _run_grading_after_claim(db: Session, task_id: int, task: HomeworkGradingTask) -> None:
    attempt = db.query(HomeworkAttempt).filter(HomeworkAttempt.id == task.attempt_id).first()
    if not attempt:
        _mark_task_failed(db, task, "attempt_not_found", "找不到对应的提交记录。")
        return
    homework = db.query(Homework).filter(Homework.id == task.homework_id).first()
    if not homework:
        _mark_task_failed(db, task, "homework_not_found", "找不到对应的作业。")
        return
    if not homework.auto_grading_enabled:
        _mark_task_failed(db, task, "auto_grading_disabled", "当前作业未启用自动评分。")
        return
    if not task.subject_id:
        _mark_task_failed(db, task, "course_missing", "当前作业未关联课程，无法使用课程级 LLM 配置。")
        return

    config = (
        db.query(CourseLLMConfig)
        .options(
            joinedload(CourseLLMConfig.groups)
            .joinedload(LLMGroup.members)
            .joinedload(CourseLLMConfigEndpoint.preset),
            joinedload(CourseLLMConfig.endpoints).joinedload(CourseLLMConfigEndpoint.preset),
        )
        .filter(CourseLLMConfig.subject_id == task.subject_id)
        .first()
    )
    if not config or not config.is_enabled:
        _mark_task_failed(db, task, "llm_config_disabled", "当前课程未启用 LLM 配置。")
        return

    if not (config.groups or []) and not (config.endpoints or []):
        _mark_task_failed(db, task, "endpoint_missing", "当前课程未配置可用端点。")
        return

    routing_warnings = _homework_routing_warnings(db, homework, config)
    material = _build_student_material(db, homework, attempt, config)
    base_manifest = material["artifact_manifest"] or {}
    if not isinstance(base_manifest, dict):
        base_manifest = {}
    task.artifact_manifest = {**base_manifest, "llm_routing": {"version": 1, "status": "pending"}}
    if routing_warnings:
        task.artifact_manifest["homework_routing_warnings"] = routing_warnings
        _flag_artifact_manifest_modified(task)
    task.input_token_estimate = material["estimated_tokens"]
    task.task_summary = material["summary"]

    if material["all_empty"]:
        skipped = base_manifest.get("skipped") if isinstance(base_manifest, dict) else None
        detail = "附件处理后没有可评分内容。"
        if isinstance(skipped, list) and skipped:
            lines = ", ".join(f"{s.get('path', '?')}（{s.get('reason', '')}）" for s in skipped[:5])
            detail = f"{detail} 未纳入：{lines}" + (" 等。" if len(skipped) > 5 else "")
        _mark_task_failed(db, task, "no_usable_content", detail)
        return

    allowed, error_code = reserve_quota_tokens(
        db,
        task,
        config,
        estimated_tokens=int(material["estimated_tokens"] or 0),
    )
    if not allowed:
        quota_msg = {
            "quota_exceeded_student": "已达到本日学生 token 上限，自动评分未执行。",
        }.get(error_code or "", "今日额度已用尽，自动评分未执行。")
        _mark_task_failed(db, task, error_code or "quota_exceeded_student", quota_msg)
        return

    teacher_exists = (
        db.query(HomeworkScoreCandidate)
        .filter(
            HomeworkScoreCandidate.attempt_id == attempt.id,
            HomeworkScoreCandidate.source == "teacher",
        )
        .first()
    )
    if teacher_exists:
        release_quota_reservation(db, task.id)
        task.status = "success"
        task.error_code = None
        task.error_message = None
        task.finished_at = _now_utc()
        task.task_summary = "已跳过：该次提交已有教师评分，未调用模型。"
        summary = (
            db.query(HomeworkSubmission)
            .filter(
                HomeworkSubmission.homework_id == attempt.homework_id,
                HomeworkSubmission.student_id == attempt.student_id,
            )
            .first()
        )
        if summary:
            summary.latest_task_status = task.status
            summary.latest_task_error = None
            refresh_submission_summary(db, summary)
        db.commit()
        return

    try:
        response = _grade_with_endpoint_group(
            db=db,
            task=task,
            homework=homework,
            attempt=attempt,
            config=config,
            material=material,
        )
    except NonRetryableLLMError as exc:
        msg = str(exc) or "LLM 调用失败。"
        _append_llm_call_log(
            task,
            {"phase": "routing_failed", "level": "error", "message": msg},
        )
        _mark_task_failed(db, task, "llm_call_failed", msg)
        return

    candidate = HomeworkScoreCandidate(
        attempt_id=attempt.id,
        homework_id=homework.id,
        student_id=attempt.student_id,
        source="auto",
        score=normalize_score_for_homework(homework, response["score"]),
        comment=response["comment"],
        source_metadata={
            "task_id": task.id,
            "endpoint_id": response["endpoint_id"],
            "raw_response_excerpt": response["raw_response"][:1000],
        },
    )
    db.add(candidate)
    db.flush()

    task.status = "success"
    task.error_code = None
    task.error_message = None
    task.finished_at = _now_utc()
    task.task_summary = "评分成功"
    record_usage_if_needed(db, task, config, response["usage"])

    summary = (
        db.query(HomeworkSubmission)
        .filter(
            HomeworkSubmission.homework_id == attempt.homework_id,
            HomeworkSubmission.student_id == attempt.student_id,
        )
        .first()
    )
    if summary:
        summary.latest_task_status = task.status
        summary.latest_task_error = None
        refresh_submission_summary(db, summary)

    db.commit()


def _maybe_queue_auto_retry(db: Session, attempt: HomeworkAttempt, error_code: str) -> None:
    """After a failed grading task (committed), optionally enqueue one more try for transient LLM errors."""
    if error_code not in _AUTO_RETRY_ELIGIBLE_ERROR_CODES:
        return
    failed_count = (
        db.query(func.count(HomeworkGradingTask.id))
        .filter(HomeworkGradingTask.attempt_id == attempt.id, HomeworkGradingTask.status == "failed")
        .scalar()
    )
    if int(failed_count or 0) > _MAX_AUTO_RETRY_TASKS_PER_ATTEMPT:
        return
    queue_grading_task(db, attempt, queue_reason=f"auto_retry_after_{error_code}")
    db.commit()


def _mark_task_failed(
    db: Session,
    task: HomeworkGradingTask,
    error_code: str,
    error_message: str,
) -> None:
    release_quota_reservation(db, task.id)
    task.status = "failed"
    task.error_code = error_code
    task.error_message = error_message
    task.finished_at = _now_utc()
    task.task_summary = error_message
    summary = (
        db.query(HomeworkSubmission)
        .filter(
            HomeworkSubmission.homework_id == task.homework_id,
            HomeworkSubmission.student_id == task.student_id,
        )
        .first()
    )
    if summary:
        summary.latest_task_status = task.status
        summary.latest_task_error = error_message
        refresh_submission_summary(db, summary)
    db.commit()
    if error_code in _AUTO_RETRY_ELIGIBLE_ERROR_CODES and task.attempt_id:
        attempt_row = db.query(HomeworkAttempt).filter(HomeworkAttempt.id == task.attempt_id).first()
        if attempt_row:
            _maybe_queue_auto_retry(db, attempt_row, error_code)


def _collect_grading_endpoints_for_config(
    config: CourseLLMConfig,
) -> tuple[list[LLMGroup], list[CourseLLMConfigEndpoint]]:
    """Return (ordered groups, flat legacy endpoints when no groups are defined)."""
    groups = sorted(
        [g for g in (config.groups or []) if g is not None],
        key=lambda x: (x.priority, x.id),
    )
    if groups and any((g.members or []) for g in groups):
        return groups, []
    flat = sorted(
        (config.endpoints or []),
        key=lambda row: (row.priority, row.id),
    )
    return [], flat


def _grade_with_endpoint_group(
    *,
    db: Session,
    task: HomeworkGradingTask,
    homework: Homework,
    attempt: HomeworkAttempt,
    config: CourseLLMConfig,
    material: dict[str, Any],
) -> dict[str, Any]:
    from types import SimpleNamespace

    need_vision = _material_needs_vision(material)
    spec = homework.llm_routing_spec if isinstance(homework.llm_routing_spec, dict) else None
    group_rows, flat_endpoints = _collect_grading_endpoints_for_config(config)

    if spec and spec.get("mode") == "latest_passing_validated":
        preset = (
            db.query(LLMEndpointPreset)
            .filter(
                LLMEndpointPreset.is_active.is_(True),
                LLMEndpointPreset.validation_status == "validated",
                LLMEndpointPreset.text_validation_status == "passed",
            )
            .order_by(LLMEndpointPreset.validated_at.desc().nullslast(), LLMEndpointPreset.id.desc())
            .first()
        )
        if not preset:
            _append_llm_call_log(
                task,
                {
                    "phase": "routing",
                    "level": "warn",
                    "message": "作业要求「最新纯文本测试通过」端点，但无可用预设，回退课程路由。",
                },
            )
        else:
            bound_ids = {m.preset_id for g in (config.groups or []) for m in (g.members or [])}
            bound_ids |= {e.preset_id for e in (config.endpoints or [])}
            if preset.id not in bound_ids:
                _append_llm_call_log(
                    task,
                    {
                        "phase": "routing",
                        "level": "warn",
                        "preset": preset.name,
                        "message": "该预设未加入本课程配置，仍将尝试直接调用。",
                    },
                )
            flat_endpoints = [
                SimpleNamespace(
                    id=-preset.id,
                    config_id=config.id,
                    group_id=None,
                    preset_id=preset.id,
                    priority=1,
                    preset=preset,
                )
            ]
            group_rows = []

    group_rows, flat_endpoints, routing_notes = _filter_course_links_for_homework(homework, group_rows, flat_endpoints)
    for note in routing_notes:
        _append_llm_call_log(task, {"phase": "routing", "level": "info", "message": note})

    if group_rows:
        routing = GroupRoutingContext.from_config(group_rows, task_id=task.id)

        def _update_routing_artifact(merge: dict[str, Any]) -> None:
            if not isinstance(task.artifact_manifest, dict):
                return
            base = dict(routing.routing_payload())
            base.update(merge)
            task.artifact_manifest["llm_routing"] = base
            _flag_artifact_manifest_modified(task)

        _update_routing_artifact({"status": "routing"})

        last_error: Optional[str] = None
        global_index = 0
        for group_state in routing.group_states:
            group_state.apply_round_robin_start(task.id)
            while group_state.current_order:
                link = group_state.current_order[0]
                global_index += 1
                preset: LLMEndpointPreset = link.preset
                ok, reason = _preset_eligible_for_grading(preset, need_vision=need_vision)
                if not ok:
                    last_error = reason
                    _append_llm_call_log(
                        task,
                        {
                            "phase": "skip_endpoint",
                            "level": "warn",
                            "preset": getattr(preset, "name", None),
                            "message": reason,
                        },
                    )
                    group_state.remove_member(link)
                    _update_routing_artifact(
                        {
                            "status": "invalid_member_skipped",
                            "last_error": last_error,
                        }
                    )
                    continue
                task.current_endpoint_index = global_index
                attempt_limit = max(1, int(preset.max_retries or 0) + 1)
                member_done = False
                for request_attempt in range(1, attempt_limit + 1):
                    task.current_attempt = request_attempt
                    try:
                        score_result = _request_grade_from_endpoint(
                            preset=preset,
                            homework=homework,
                            attempt=attempt,
                            config=config,
                            material=material,
                            task=task,
                        )
                        score_result["endpoint_id"] = preset.id
                        _update_routing_artifact({"status": "ok"})
                        return score_result
                    except RetryableLLMError as exc:
                        last_error = str(exc)
                        n_before = len(group_state.current_order)
                        routing.note_failure(group_state, link, exc)
                        if request_attempt >= attempt_limit:
                            if n_before == 1:
                                group_state.remove_member(link)
                            _update_routing_artifact(
                                {
                                    "status": "adaptive_shift",
                                    "last_error": last_error,
                                }
                            )
                            member_done = True
                            break
                        _update_routing_artifact(
                            {
                                "status": "retry_backoff",
                                "last_error": last_error,
                            }
                        )
                        wait_seconds = min(
                            int(preset.initial_backoff_seconds or 2) * (2 ** (request_attempt - 1)),
                            120,
                        )
                        if os.environ.get("LLM_GRADING_TEST_SKIP_BACKOFF") != "1":
                            time.sleep(wait_seconds)
                    except NonRetryableLLMError as exc:
                        last_error = str(exc)
                        routing.note_failure(group_state, link, exc)
                        _update_routing_artifact({"status": "endpoint_error", "last_error": last_error})
                        group_state.remove_member(link)
                        member_done = True
                        break
        _update_routing_artifact({"status": "failed", "message": last_error or ""})
        raise NonRetryableLLMError(last_error or "所有组内端点都调用失败。")

    last_error_flat: Optional[str] = None
    for endpoint_index, link in enumerate(flat_endpoints, start=1):
        preset: LLMEndpointPreset = link.preset
        ok, reason = _preset_eligible_for_grading(preset, need_vision=need_vision)
        if not ok:
            last_error_flat = reason
            _append_llm_call_log(
                task,
                {
                    "phase": "skip_endpoint",
                    "level": "warn",
                    "preset": getattr(preset, "name", None),
                    "message": reason,
                },
            )
            continue
        task.current_endpoint_index = endpoint_index
        attempt_limit = max(1, int(preset.max_retries or 0) + 1)
        for request_attempt in range(1, attempt_limit + 1):
            task.current_attempt = request_attempt
            try:
                score_result = _request_grade_from_endpoint(
                    preset=preset,
                    homework=homework,
                    attempt=attempt,
                    config=config,
                    material=material,
                    task=task,
                )
                score_result["endpoint_id"] = preset.id
                if isinstance(task.artifact_manifest, dict) and "llm_routing" in (task.artifact_manifest or {}):
                    task.artifact_manifest["llm_routing"] = (task.artifact_manifest.get("llm_routing") or {}) | {
                        "version": 1,
                        "mode": "legacy_priority",
                        "status": "ok",
                    }
                    _flag_artifact_manifest_modified(task)
                return score_result
            except RetryableLLMError as exc:
                last_error_flat = str(exc)
                if request_attempt >= attempt_limit:
                    break
                wait_seconds = min(
                    int(preset.initial_backoff_seconds or 2) * (2 ** (request_attempt - 1)),
                    120,
                )
                if os.environ.get("LLM_GRADING_TEST_SKIP_BACKOFF") != "1":
                    time.sleep(wait_seconds)
            except NonRetryableLLMError as exc:
                last_error_flat = str(exc)
                break
    if isinstance(task.artifact_manifest, dict) and "llm_routing" in (task.artifact_manifest or {}):
        task.artifact_manifest["llm_routing"] = (task.artifact_manifest.get("llm_routing") or {}) | {
            "version": 1,
            "mode": "legacy_priority",
            "status": "failed",
        }
        _flag_artifact_manifest_modified(task)
    raise NonRetryableLLMError(last_error_flat or "所有端点都调用失败。")


def _build_chat_completion_url(base_url: str) -> str:
    normalized = base_url.rstrip("/") + "/"
    if normalized.endswith("/chat/completions/"):
        return normalized[:-1]
    return urljoin(normalized, "chat/completions")


def _redact_endpoint_host(url: str) -> str:
    try:
        parsed = urlparse(url)
        if parsed.scheme and parsed.netloc:
            return f"{parsed.scheme}://{parsed.netloc}/…/chat/completions"
    except Exception:
        pass
    return "（端点 URL 已省略）"


def _request_grade_from_endpoint(
    *,
    preset: LLMEndpointPreset,
    homework: Homework,
    attempt: HomeworkAttempt,
    config: CourseLLMConfig,
    material: dict[str, Any],
    task: Optional[HomeworkGradingTask] = None,
) -> dict[str, Any]:
    timeout = httpx.Timeout(
        connect=preset.connect_timeout_seconds or 10,
        read=preset.read_timeout_seconds or 120,
        write=preset.read_timeout_seconds or 120,
        pool=preset.connect_timeout_seconds or 10,
    )
    payload = {
        "model": preset.model_name,
        "messages": _build_scoring_messages(homework, attempt, config, material),
        "temperature": 0.2,
        "max_tokens": config.max_output_tokens or 1200,
    }
    endpoint_url = _build_chat_completion_url(preset.base_url)
    if task is not None:
        _append_llm_call_log(
            task,
            {
                "phase": "http_request",
                "level": "info",
                "preset": preset.name,
                "model": preset.model_name,
                "endpoint": _redact_endpoint_host(endpoint_url),
            },
        )
    try:
        with httpx.Client(timeout=timeout) as client:
            response = client.post(
                endpoint_url,
                json=payload,
                headers={
                    "Authorization": f"Bearer {preset.api_key}",
                    "Content-Type": "application/json",
                },
            )
    except httpx.TimeoutException as exc:
        if task is not None:
            _append_llm_call_log(
                task,
                {"phase": "http_error", "level": "error", "preset": preset.name, "message": f"请求超时：{exc}"},
            )
        raise RetryableLLMError(f"请求超时：{exc}") from exc
    except httpx.HTTPError as exc:
        if task is not None:
            _append_llm_call_log(
                task,
                {"phase": "http_error", "level": "error", "preset": preset.name, "message": f"网络请求失败：{exc}"},
            )
        raise RetryableLLMError(f"网络请求失败：{exc}") from exc

    if response.status_code in NON_RETRYABLE_STATUS_CODES:
        if task is not None:
            _append_llm_call_log(
                task,
                {
                    "phase": "http_response",
                    "level": "error",
                    "preset": preset.name,
                    "http_status": response.status_code,
                    "body_excerpt": (response.text or "")[:400],
                },
            )
        raise NonRetryableLLMError(f"鉴权或权限失败：HTTP {response.status_code}")
    if response.status_code == 413:
        if task is not None:
            _append_llm_call_log(
                task,
                {"phase": "http_response", "level": "error", "preset": preset.name, "http_status": 413},
            )
        raise NonRetryableLLMError("请求内容过大，端点拒绝处理。")
    if response.status_code in RETRYABLE_STATUS_CODES:
        if task is not None:
            _append_llm_call_log(
                task,
                {
                    "phase": "http_response",
                    "level": "warn",
                    "preset": preset.name,
                    "http_status": response.status_code,
                    "body_excerpt": (response.text or "")[:400],
                },
            )
        raise RetryableLLMError(f"端点暂时不可用：HTTP {response.status_code}")
    if response.status_code >= 400:
        if task is not None:
            _append_llm_call_log(
                task,
                {
                    "phase": "http_response",
                    "level": "error",
                    "preset": preset.name,
                    "http_status": response.status_code,
                    "body_excerpt": (response.text or "")[:400],
                },
            )
        raise NonRetryableLLMError(f"端点请求失败：HTTP {response.status_code} {response.text[:300]}")

    try:
        data = response.json()
    except ValueError as exc:
        if task is not None:
            _append_llm_call_log(
                task,
                {"phase": "parse_error", "level": "warn", "preset": preset.name, "message": "模型返回的不是合法 JSON 响应。"},
            )
        raise RetryableLLMError("模型返回的不是合法 JSON 响应。") from exc

    raw_content = _extract_message_content(data)
    try:
        score_payload = _parse_scoring_json(raw_content, homework)
    except RetryableLLMError as exc:
        if task is not None:
            _append_llm_call_log(
                task,
                {
                    "phase": "parse_model_output",
                    "level": "warn",
                    "preset": preset.name,
                    "message": str(exc),
                    "raw_excerpt": (raw_content or "")[:500],
                },
            )
        raise
    usage = data.get("usage") or {}
    if task is not None:
        _append_llm_call_log(
            task,
            {
                "phase": "success",
                "level": "info",
                "preset": preset.name,
                "usage": usage,
            },
        )
    return {
        "score": score_payload["score"],
        "comment": score_payload["comment"],
        "usage": {
            "prompt_tokens": usage.get("prompt_tokens"),
            "completion_tokens": usage.get("completion_tokens"),
            "total_tokens": usage.get("total_tokens"),
        },
        "raw_response": raw_content,
    }


def _extract_message_content(response_json: dict[str, Any]) -> str:
    choices = response_json.get("choices") or []
    if not choices:
        return ""
    message = choices[0].get("message") or {}
    content = message.get("content")
    if isinstance(content, str):
        return content
    if isinstance(content, list):
        parts: list[str] = []
        for item in content:
            if isinstance(item, str):
                parts.append(item)
            elif isinstance(item, dict) and item.get("type") == "text":
                parts.append(item.get("text") or "")
        return "\n".join(part for part in parts if part)
    return ""


def _strip_markdown_fence(text: str) -> str:
    match = JSON_FENCE_PATTERN.match(text or "")
    if match:
        return match.group(1).strip()
    return (text or "").strip()


def _extract_first_json_object(text: str) -> Optional[str]:
    start = text.find("{")
    if start < 0:
        return None
    depth = 0
    in_string = False
    escaped = False
    for index, char in enumerate(text[start:], start=start):
        if in_string:
            if escaped:
                escaped = False
            elif char == "\\":
                escaped = True
            elif char == '"':
                in_string = False
            continue
        if char == '"':
            in_string = True
            continue
        if char == "{":
            depth += 1
        elif char == "}":
            depth -= 1
            if depth == 0:
                return text[start : index + 1]
    return None


def _parse_scoring_json(raw_text: str, homework: Homework) -> dict[str, Any]:
    text = _strip_markdown_fence(raw_text)
    payload_text = text
    try:
        payload = json.loads(payload_text)
    except ValueError:
        payload_text = _extract_first_json_object(text or "") or ""
        if not payload_text:
            raise RetryableLLMError("模型未按要求输出 JSON 对象。")
        try:
            payload = json.loads(payload_text)
        except ValueError as exc:
            raise RetryableLLMError("模型输出 JSON 解析失败。") from exc
    if not isinstance(payload, dict):
        raise RetryableLLMError("模型输出的根节点不是对象。")
    if "score" not in payload or "comment" not in payload:
        raise RetryableLLMError("模型输出缺少 score/comment 字段。")
    try:
        score = float(payload["score"])
    except (TypeError, ValueError) as exc:
        raise RetryableLLMError("模型输出的 score 不是有效数字。") from exc
    if score < 0 or score > float(homework.max_score or 100):
        raise RetryableLLMError("模型输出分数超出允许范围。")
    comment = payload.get("comment")
    if comment is None:
        comment = ""
    if not isinstance(comment, str):
        comment = str(comment)
    return {"score": score, "comment": comment}


def _build_scoring_messages(
    homework: Homework,
    attempt: HomeworkAttempt,
    config: CourseLLMConfig,
    material: dict[str, Any],
) -> list[dict[str, Any]]:
    system_prompt = config.system_prompt or (
        "你是一个严格遵守格式要求的课程作业评分助手。"
        "你必须只输出 JSON 对象，且字段固定为 score 与 comment。"
        "绝不能在 JSON 前后输出任何额外说明。"
    )
    response_language = homework.response_language or config.response_language or "zh-CN"
    teacher_prompt = config.teacher_prompt or ""
    assignment_body = "\n\n".join(material["assignment_texts"])
    assignment_text = f"{SECTION_ASSIGNMENT}\n### 教师侧作业说明与材料\n{assignment_body}"
    if teacher_prompt.strip():
        assignment_text += f"\n\n### 教师补充提示（课程 LLM 配置）\n{teacher_prompt}"
    student_intro = (
        f"{SECTION_STUDENT_BODY}\n### 提交元数据\n"
        f"作业标题：{homework.title}\n"
        f"满分：{normalize_score_for_homework(homework, homework.max_score)}\n"
        f"评分精度：{'1 位小数' if homework.grade_precision == 'decimal_1' else '整数'}\n"
        f"响应语言：{response_language}\n"
        f"提交是否迟交：{'是' if attempt.is_late else '否'}\n"
        f"迟交默认是否影响得分：{'是' if homework.late_submission_affects_score else '否'}\n"
    )
    user_parts: list[dict[str, Any]] = [{"type": "text", "text": assignment_text}]
    user_parts.append({"type": "text", "text": student_intro})
    text_blocks = [b for b in material["student_blocks"] if b.block_type == "text"]
    image_blocks = [b for b in material["student_blocks"] if b.block_type == "image"]
    if text_blocks:
        user_parts.append(
            {
                "type": "text",
                "text": f"{SECTION_ATTACHMENT}\n（以下为学生在表单中的说明文字，以及从附件解析出的可读文本）",
            }
        )
        for block in text_blocks:
            meta = _attachment_block_meta_text(block)
            user_parts.append({"type": "text", "text": (meta + (block.text or "")).strip()})
    if image_blocks:
        user_parts.append({"type": "text", "text": f"{SECTION_IMAGES}\n（以下为提交中的图片/PDF 页渲染，按顺序评分）"})
        for block in image_blocks:
            cap = (
                f"[IMAGE_META path={block.logical_path or block.path} "
                f"mime={block.mime_hint or 'image'} origin={block.origin or 'attachment'}]"
            )
            user_parts.append({"type": "text", "text": cap})
            user_parts.append({"type": "image_url", "image_url": {"url": block.image_data_url}})
    if material["notes_text"]:
        user_parts.append({"type": "text", "text": f"{SECTION_NOTES}\n{material['notes_text']}"})
    return [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_parts},
    ]


def _truncate_text(value: str, limit: int = MAX_FILE_TEXT_CHARS) -> tuple[str, bool]:
    if len(value) <= limit:
        return value, False
    return value[:limit], True


def _decode_bytes_as_text(content: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gbk", "gb2312", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="ignore")


def _safe_relative_path(path_text: str) -> Optional[str]:
    normalized = PurePosixPath(path_text.replace("\\", "/"))
    safe_parts = []
    for part in normalized.parts:
        if part in {"", ".", ".."}:
            continue
        safe_parts.append(part)
    if not safe_parts:
        return None
    return "/".join(safe_parts)


def _guess_mime_type(path: str) -> str:
    mime_type, _ = mimetypes.guess_type(path)
    return mime_type or "application/octet-stream"


def _bytes_to_data_url(path: str, content: bytes) -> str:
    mime_type = _guess_mime_type(path)
    return f"data:{mime_type};base64,{base64.b64encode(content).decode('ascii')}"


def _extract_docx_text(content: bytes) -> str:
    document = Document(io.BytesIO(content))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        t = (paragraph.text or "").strip()
        if t:
            parts.append(t)
    for table in document.tables:
        rows_out: list[str] = []
        for row in table.rows:
            cells = [" ".join((c.text or "").split()) for c in row.cells]
            if any(cells):
                rows_out.append("\t".join(cells))
        if rows_out:
            parts.append("[表格]\n" + "\n".join(rows_out))
    return "\n".join(parts)


def _extract_xlsx_text(content: bytes, path: str) -> str:
    wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    try:
        parts: list[str] = []
        max_rows = 300
        max_cols = 64
        for sheet in wb.worksheets:
            parts.append(f"[工作表: {sheet.title}]")
            for row in sheet.iter_rows(min_row=1, max_row=max_rows, max_col=max_cols, values_only=True):
                cells = ["" if v is None else str(v).strip() for v in row]
                if any(cells):
                    parts.append("\t".join(cells))
        return "\n".join(parts).strip()
    finally:
        wb.close()


def _extract_xls_text(content: bytes, path: str) -> str:
    book = xlrd.open_workbook(file_contents=content, formatting_info=False, on_demand=True)
    parts: list[str] = []
    max_rows = 300
    max_cols = 64
    try:
        for sheet in book.sheets():
            parts.append(f"[工作表: {sheet.name}]")
            for rx in range(min(sheet.nrows, max_rows)):
                row_vals = []
                for cx in range(min(sheet.ncols, max_cols)):
                    row_vals.append(str(sheet.cell_value(rx, cx)).strip())
                if any(row_vals):
                    parts.append("\t".join(row_vals))
    finally:
        book.release_resources()
    return "\n".join(parts).strip()


def _extract_legacy_doc_text(content: bytes) -> str:
    """
    Best-effort text from legacy .doc (OLE compound file). Does not preserve layout/tables.
    Reads the main Word binary streams only (avoids pulling unrelated OLE blobs).
    """
    if not olefile.isOleFile(io.BytesIO(content)):
        return ""
    ole = olefile.OleFileIO(io.BytesIO(content))
    max_stream_bytes = 6 * 1024 * 1024
    blob = b""
    try:
        for s in ole.listdir():
            joined = "/".join(s).lower()
            if not (
                "worddocument" in joined
                or joined.endswith("1table")
                or joined.endswith("0table")
                or joined == "data"
            ):
                continue
            try:
                blob += ole.openstream(s).read(max_stream_bytes)
            except Exception:
                continue
        if not blob:
            for s in ole.listdir():
                try:
                    blob += ole.openstream(s).read(min(max_stream_bytes, 512 * 1024))
                    if len(blob) >= max_stream_bytes:
                        break
                except Exception:
                    continue
    finally:
        ole.close()
    if not blob:
        return ""
    text = blob.decode("utf-16le", errors="ignore")
    lines: list[str] = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if len(line) < 2:
            continue
        printable = sum(1 for c in line if c.isprintable() or c in "\t")
        if printable / max(len(line), 1) < 0.55:
            continue
        lines.append(line)
    return "\n".join(lines).strip()


def _extract_pdf_images(content: bytes, path: str) -> list[MaterialBlock]:
    blocks: list[MaterialBlock] = []
    document = fitz.open(stream=content, filetype="pdf")
    try:
        for index, page in enumerate(document, start=1):
            pixmap = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
            image_bytes = pixmap.tobytes("png")
            logical = f"{path}#page-{index}"
            blocks.append(
                MaterialBlock(
                    priority=3,
                    path=logical,
                    block_type="image",
                    image_data_url=f"data:image/png;base64,{base64.b64encode(image_bytes).decode('ascii')}",
                    estimated_tokens=settings.DEFAULT_ESTIMATED_IMAGE_TOKENS,
                    logical_path=logical,
                    mime_hint="image/png",
                    origin="attachment",
                )
            )
    finally:
        document.close()
    return blocks


def _extract_ipynb_blocks(content: bytes, path: str) -> list[MaterialBlock]:
    try:
        notebook = json.loads(content.decode("utf-8"))
    except Exception:
        return []

    blocks: list[MaterialBlock] = []
    text_fragments: list[str] = []
    for index, cell in enumerate(notebook.get("cells") or [], start=1):
        cell_type = cell.get("cell_type") or "cell"
        source = "".join(cell.get("source") or [])
        if source.strip():
            text_fragments.append(f"## Cell {index} ({cell_type})\n{source.strip()}")
        for output in cell.get("outputs") or []:
            if output.get("output_type") == "stream":
                text = "".join(output.get("text") or [])
                if text.strip():
                    text_fragments.append(f"Output:\n{text.strip()[:MAX_IPYNB_OUTPUT_CHARS]}")
            data = output.get("data") or {}
            text_plain = data.get("text/plain")
            if text_plain:
                text = "".join(text_plain if isinstance(text_plain, list) else [str(text_plain)])
                if text.strip():
                    text_fragments.append(f"Output:\n{text.strip()[:MAX_IPYNB_OUTPUT_CHARS]}")
            image_png = data.get("image/png")
            if image_png:
                if isinstance(image_png, list):
                    image_png = "".join(image_png)
                logical_img = f"{path}#cell-{index}-output"
                blocks.append(
                    MaterialBlock(
                        priority=3,
                        path=logical_img,
                        block_type="image",
                        image_data_url=f"data:image/png;base64,{image_png}",
                        estimated_tokens=settings.DEFAULT_ESTIMATED_IMAGE_TOKENS,
                        logical_path=logical_img,
                        mime_hint="image/png",
                        origin="attachment",
                    )
                )
    if text_fragments:
        text, truncated = _truncate_text("\n\n".join(text_fragments))
        suffix = "\n\n[说明] Ipynb 文本输出已截断。" if truncated else ""
        mime_nb = "application/x-ipynb+json"
        blocks.insert(
            0,
            MaterialBlock(
                priority=2,
                path=path,
                block_type="text",
                text=f"### 附件（Jupyter 解析）\n**文件**: {path}\n**类型**: {mime_nb}\n\n{text}{suffix}",
                estimated_tokens=int(len(text) / 4) + 50,
                logical_path=path,
                mime_hint=mime_nb,
                origin="attachment",
                truncated=truncated,
            ),
        )
    return blocks


def _classify_and_extract(path: str, content: bytes) -> list[MaterialBlock]:
    suffix = PurePosixPath(path).suffix.lower()
    if suffix in PDF_EXTENSIONS:
        return _extract_pdf_images(content, path)
    if suffix in IPYNB_EXTENSIONS:
        return _extract_ipynb_blocks(content, path)
    if suffix in SUPPORTED_IMAGE_EXTENSIONS:
        mime_img = _guess_mime_type(path)
        return [
            MaterialBlock(
                priority=3,
                path=path,
                block_type="image",
                image_data_url=_bytes_to_data_url(path, content),
                estimated_tokens=settings.DEFAULT_ESTIMATED_IMAGE_TOKENS,
                logical_path=path,
                mime_hint=mime_img,
                origin="attachment",
            )
        ]
    if suffix == ".docx":
        text = _extract_docx_text(content)
    elif suffix in EXCEL_XLSX_EXTENSIONS:
        text = _extract_xlsx_text(content, path)
    elif suffix in EXCEL_XLS_EXTENSIONS:
        text = _extract_xls_text(content, path)
    elif suffix in LEGACY_WORD_DOC_EXTENSIONS:
        text = _extract_legacy_doc_text(content)
    elif suffix in SUPPORTED_TEXT_EXTENSIONS:
        text = _decode_bytes_as_text(content)
    else:
        return []
    text = text.strip()
    if not text:
        return []
    text, truncated = _truncate_text(text)
    suffix_note = "\n\n[说明] 文件内容过长，已截断。" if truncated else ""
    mime_doc = _guess_mime_type(path)
    return [
        MaterialBlock(
            priority=2,
            path=path,
            block_type="text",
            text=f"### 附件（解析文本）\n**文件**: {path}\n**类型**: {mime_doc}\n\n{text}{suffix_note}",
            estimated_tokens=int(len(text) / 4) + 50,
            logical_path=path,
            mime_hint=mime_doc,
            origin="attachment",
            truncated=truncated,
        )
    ]


def _walk_zip_bytes(
    content: bytes,
    *,
    root_path: str,
    depth: int,
    state: dict[str, Any],
) -> list[MaterialBlock]:
    blocks: list[MaterialBlock] = []
    if depth > MAX_ZIP_DEPTH:
        state["skipped"].append({"path": root_path, "reason": "超过最大嵌套深度"})
        return blocks
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            infos = sorted(archive.infolist(), key=lambda item: (item.filename or "").lower())
            for info in infos:
                if info.is_dir():
                    continue
                safe_child_path = _safe_relative_path(info.filename or "")
                if not safe_child_path:
                    state["skipped"].append({"path": f"{root_path}/{info.filename}", "reason": "非法路径"})
                    continue
                state["file_count"] += 1
                state["total_bytes"] += max(0, int(info.file_size or 0))
                if state["file_count"] > MAX_ZIP_FILES or state["total_bytes"] > MAX_ZIP_TOTAL_BYTES:
                    state["skipped"].append({"path": f"{root_path}/{safe_child_path}", "reason": "超出展开文件数或总大小限制"})
                    continue
                child_bytes = archive.read(info)
                child_path = f"{root_path}/{safe_child_path}"
                child_suffix = PurePosixPath(safe_child_path).suffix.lower()
                if child_suffix in ZIP_EXTENSIONS:
                    blocks.extend(_walk_zip_bytes(child_bytes, root_path=child_path, depth=depth + 1, state=state))
                elif child_suffix in RAR_EXTENSIONS:
                    blocks.extend(_walk_rar_bytes(child_bytes, root_path=child_path, depth=depth + 1, state=state))
                else:
                    child_blocks = _classify_and_extract(child_path, child_bytes)
                    if child_blocks:
                        blocks.extend(child_blocks)
                    else:
                        state["skipped"].append({"path": child_path, "reason": "无法识别或提取为空"})
    except zipfile.BadZipFile:
        state["skipped"].append({"path": root_path, "reason": "压缩包损坏或格式不合法"})
    return blocks


def _walk_rar_bytes(
    content: bytes,
    *,
    root_path: str,
    depth: int,
    state: dict[str, Any],
) -> list[MaterialBlock]:
    """
    RAR listing via rarfile; extraction via 7z stdout (RAR5 compatible, no unrar stub needed).
    """
    blocks: list[MaterialBlock] = []
    if depth > MAX_ZIP_DEPTH:
        state["skipped"].append({"path": root_path, "reason": "超过最大嵌套深度"})
        return blocks
    if not _unrar_tool_path():
        state["skipped"].append({"path": root_path, "reason": "RAR 解压需要安装 unrar 或 unrar-free"})
        return blocks
    tmp_path: Optional[str] = None
    try:
        fd, tmp_path = tempfile.mkstemp(suffix=".rar")
        os.write(fd, content)
        os.close(fd)
        with rarfile.RarFile(tmp_path) as archive:
            if archive.needs_password():
                state["skipped"].append({"path": root_path, "reason": "RAR 已加密，不支持解压评分"})
                return blocks
            infos = sorted(archive.infolist(), key=lambda item: (item.filename or "").lower())
            for info in infos:
                if info.isdir():
                    continue
                if getattr(info, "needs_password", lambda: False)():
                    state["skipped"].append({"path": root_path, "reason": "RAR 内含加密条目，不支持"})
                    return blocks
                member_name = info.filename or ""
                safe_child_path = _safe_relative_path(member_name)
                if not safe_child_path:
                    state["skipped"].append({"path": f"{root_path}/{member_name}", "reason": "非法路径"})
                    continue
                state["file_count"] += 1
                state["total_bytes"] += max(0, int(info.file_size or 0))
                if state["file_count"] > MAX_ZIP_FILES or state["total_bytes"] > MAX_ZIP_TOTAL_BYTES:
                    state["skipped"].append({"path": f"{root_path}/{safe_child_path}", "reason": "超出展开文件数或总大小限制"})
                    continue
                child_path = f"{root_path}/{safe_child_path}"
                try:
                    child_bytes = _rar_read_member_bytes(tmp_path, member_name)
                except RuntimeError as exc:
                    state["skipped"].append({"path": child_path, "reason": str(exc)[:400]})
                    return blocks
                child_suffix = PurePosixPath(safe_child_path).suffix.lower()
                if child_suffix in ZIP_EXTENSIONS:
                    blocks.extend(_walk_zip_bytes(child_bytes, root_path=child_path, depth=depth + 1, state=state))
                elif child_suffix in RAR_EXTENSIONS:
                    blocks.extend(_walk_rar_bytes(child_bytes, root_path=child_path, depth=depth + 1, state=state))
                else:
                    child_blocks = _classify_and_extract(child_path, child_bytes)
                    if child_blocks:
                        blocks.extend(child_blocks)
                    else:
                        state["skipped"].append({"path": child_path, "reason": "无法识别或提取为空"})
    except rarfile.BadRarFile:
        state["skipped"].append({"path": root_path, "reason": "RAR 损坏或格式不合法"})
    except rarfile.NotRarFile:
        state["skipped"].append({"path": root_path, "reason": "不是有效的 RAR 文件"})
    finally:
        if tmp_path:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass
    return blocks


def _collect_attachment_blocks(summary_path: str, attachment_name: str) -> tuple[list[MaterialBlock], list[dict[str, str]]]:
    file_path = get_attachment_file_path(summary_path)
    if not file_path or not file_path.exists():
        return [], [{"path": attachment_name or "attachment", "reason": "找不到原始附件文件"}]
    content = file_path.read_bytes()
    suffix = file_path.suffix.lower()
    state = {"file_count": 0, "total_bytes": 0, "skipped": []}
    name_lower = (attachment_name or "").lower()
    if suffix in ZIP_EXTENSIONS or name_lower.endswith(".zip"):
        blocks = _walk_zip_bytes(content, root_path=attachment_name or file_path.name, depth=1, state=state)
        return blocks, state["skipped"]
    if suffix in RAR_EXTENSIONS or name_lower.endswith(".rar"):
        blocks = _walk_rar_bytes(content, root_path=attachment_name or file_path.name, depth=1, state=state)
        return blocks, state["skipped"]
    blocks = _classify_and_extract(attachment_name or file_path.name, content)
    if blocks:
        return blocks, []
    return [], [{"path": attachment_name or file_path.name, "reason": "无法识别或提取为空"}]


def _best_score_candidate_for_attempt(db: Session, attempt_id: int) -> Optional[HomeworkScoreCandidate]:
    """Pick the display score row for one attempt (same rule as teacher submission history)."""
    candidates = (
        db.query(HomeworkScoreCandidate)
        .filter(HomeworkScoreCandidate.attempt_id == attempt_id)
        .order_by(HomeworkScoreCandidate.updated_at.desc(), HomeworkScoreCandidate.id.desc())
        .all()
    )
    if not candidates:
        return None
    return max(
        candidates,
        key=lambda item: (
            float(item.score or 0),
            1 if item.source == "teacher" else 0,
            item.updated_at or item.created_at,
        ),
    )


def _format_iteration_context_for_prompt(db: Session, homework: Homework, attempt: HomeworkAttempt) -> Optional[str]:
    """
    Text-only summary of the last N prior attempts for the same student+homework (same submission summary).
    Full multimodal re-parse of old attachments is intentionally omitted to save tokens.
    """
    summary_id = attempt.submission_summary_id
    if not summary_id:
        return None
    priors = (
        db.query(HomeworkAttempt)
        .filter(
            HomeworkAttempt.homework_id == homework.id,
            HomeworkAttempt.student_id == attempt.student_id,
            HomeworkAttempt.submission_summary_id == summary_id,
            HomeworkAttempt.id < attempt.id,
        )
        .order_by(HomeworkAttempt.id.desc())
        .limit(ITERATION_CONTEXT_MAX_PRIOR_ATTEMPTS)
        .all()
    )
    if not priors:
        return None
    priors_chrono = list(reversed(priors))
    lines: list[str] = [
        "### 迭代上下文（仅保留最近 "
        f"{ITERATION_CONTEXT_MAX_PRIOR_ATTEMPTS} 次历史提交的文字摘要；更早轮次已省略以节省 token）",
        "以下为该学生此前提交的要点，供你判断是否在反馈基础上有改进（当前要评的是最新一次提交，见后文）。",
    ]
    for idx, prev in enumerate(priors_chrono, start=1):
        cand = _best_score_candidate_for_attempt(db, prev.id)
        score_part = ""
        if cand is not None and cand.score is not None:
            src = "教师" if cand.source == "teacher" else "自动"
            score_part = f"当时展示分（{src}）：{normalize_score_for_homework(homework, cand.score)}。"
        comment = (cand.comment or "").strip() if cand else ""
        if len(comment) > ITERATION_PRIOR_COMMENT_CHAR_MAX:
            comment = comment[:ITERATION_PRIOR_COMMENT_CHAR_MAX] + "…"
        body = (prev.content or "").strip()
        if len(body) > ITERATION_PRIOR_NOTE_CHAR_MAX:
            body = body[:ITERATION_PRIOR_NOTE_CHAR_MAX] + "…"
        att = "有附件" if prev.attachment_url else "无附件"
        att_name = f"（{prev.attachment_name}）" if prev.attachment_name else ""
        lines.append(f"- 历史第 {idx} 轮：{att}{att_name}。{score_part}")
        if body:
            lines.append(f"  学生说明摘录：{body}")
        if comment:
            lines.append(f"  当时评语摘录：{comment}")
    lines.append(
        "请结合上述有限历史与当前稿评分；若当前稿明显回应了此前评语中的问题，可在 score 上合理体现进步。"
    )
    return "\n".join(lines)


def _build_student_material(
    db: Session,
    homework: Homework,
    attempt: HomeworkAttempt,
    config: CourseLLMConfig,
) -> dict[str, Any]:
    assignment_texts = [
        f"作业标题：{homework.title}",
        f"作业要求：\n{homework.content or '无'}",
    ]
    iteration_ctx = _format_iteration_context_for_prompt(db, homework, attempt)
    if iteration_ctx:
        assignment_texts.append(iteration_ctx)
    if homework.reference_answer:
        assignment_texts.append(f"参考答案或提示：\n{homework.reference_answer}")
    if homework.rubric_text:
        assignment_texts.append(f"评分要点：\n{homework.rubric_text}")

    student_blocks: list[MaterialBlock] = []
    skipped: list[dict[str, str]] = []
    if attempt.content:
        text, truncated = _truncate_text(attempt.content)
        note = "\n\n[说明] 提交说明过长，已截断。" if truncated else ""
        student_blocks.append(
            MaterialBlock(
                priority=2,
                path="submission-note",
                block_type="text",
                text=f"### 学生正文（提交说明）\n{text}{note}",
                estimated_tokens=int(len(text) / 4) + 50,
                logical_path="submission-note",
                mime_hint="text/plain",
                origin="submission_body",
                truncated=truncated,
            )
        )
    if attempt.attachment_url:
        attachment_blocks, skipped_items = _collect_attachment_blocks(
            attempt.attachment_url,
            attempt.attachment_name or "attachment",
        )
        student_blocks.extend(attachment_blocks)
        skipped.extend(skipped_items)

    student_blocks.sort(key=lambda item: (item.priority, item.path))
    text_budget = int((config.max_input_tokens or 16000) * (config.estimated_chars_per_token or 4.0))
    reserved_text = "\n\n".join(assignment_texts)
    remaining_chars = max(2000, text_budget - len(reserved_text))
    remaining_image_budget = config.max_input_tokens or 16000
    final_blocks: list[MaterialBlock] = []
    truncation_notes: list[str] = []
    for block in student_blocks:
        if block.block_type == "text":
            block_text = block.text or ""
            if remaining_chars <= 0:
                skipped.append({"path": block.path, "reason": "超出输入长度预算"})
                continue
            if len(block_text) > remaining_chars:
                truncated_text, _ = _truncate_text(block_text, remaining_chars)
                final_blocks.append(
                    MaterialBlock(
                        priority=block.priority,
                        path=block.path,
                        block_type="text",
                        text=truncated_text,
                        estimated_tokens=int(len(truncated_text) / 4) + 50,
                        logical_path=block.logical_path,
                        mime_hint=block.mime_hint,
                        origin=block.origin,
                        truncated=True,
                    )
                )
                truncation_notes.append(f"{block.path} 已按预算截断")
                remaining_chars = 0
                continue
            final_blocks.append(block)
            remaining_chars -= len(block_text)
        else:
            estimated_tokens = block.estimated_tokens or settings.DEFAULT_ESTIMATED_IMAGE_TOKENS
            if remaining_image_budget < estimated_tokens:
                skipped.append({"path": block.path, "reason": "超出图片 token 预算"})
                continue
            final_blocks.append(block)
            remaining_image_budget -= estimated_tokens

    notes_text_parts: list[str] = []
    if truncation_notes:
        notes_text_parts.append("截断说明：\n- " + "\n- ".join(truncation_notes))
    if skipped:
        skipped_lines = [f"{item['path']}：{item['reason']}" for item in skipped]
        notes_text_parts.append("未纳入内容：\n- " + "\n- ".join(skipped_lines))
    notes_text = "\n\n".join(notes_text_parts)
    teacher_prompt = config.teacher_prompt or ""
    assignment_joined = "\n\n".join(assignment_texts)
    student_intro = (
        f"作业标题：{homework.title}\n"
        f"满分：{normalize_score_for_homework(homework, homework.max_score)}\n"
        f"评分精度：{'1 位小数' if homework.grade_precision == 'decimal_1' else '整数'}\n"
        f"响应语言：{homework.response_language or config.response_language or 'zh-CN'}\n"
        f"提交是否迟交：{'是' if attempt.is_late else '否'}\n"
        f"迟交默认是否影响得分：{'是' if homework.late_submission_affects_score else '否'}\n"
    )
    temp_material = {
        "student_blocks": final_blocks,
        "notes_text": notes_text,
    }
    estimated_tokens = estimate_request_tokens_from_material(
        config,
        temp_material,
        assignment_text=f"{SECTION_ASSIGNMENT}\n{assignment_joined}",
        teacher_prompt=teacher_prompt,
        student_intro=student_intro,
    )
    artifact_manifest = {
        "included": [
            {
                "path": block.path,
                "type": block.block_type,
                "logical_path": block.logical_path,
                "mime_hint": block.mime_hint,
                "origin": block.origin,
                "truncated": block.truncated,
            }
            for block in final_blocks
        ],
        "skipped": skipped,
    }
    return {
        "assignment_texts": assignment_texts,
        "student_blocks": final_blocks,
        "notes_text": notes_text,
        "estimated_tokens": estimated_tokens,
        "artifact_manifest": artifact_manifest,
        "summary": f"纳入 {len(final_blocks)} 个片段，跳过 {len(skipped)} 个文件/片段",
        "all_empty": len(final_blocks) == 0,
    }
