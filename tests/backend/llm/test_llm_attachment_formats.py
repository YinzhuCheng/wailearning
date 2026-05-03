"""Tests for LLM grading attachment parsing: Office formats and RAR."""

from __future__ import annotations

import io
import shutil
import subprocess
import zipfile

import pytest
from docx import Document

from apps.backend.wailearning_backend.llm_grading import _classify_and_extract, _walk_rar_bytes, _walk_zip_bytes


def _blocks_to_text(blocks: list) -> str:
    return "\n".join((b.text or "") for b in blocks if getattr(b, "block_type", None) == "text")


def test_extract_docx_includes_table_cells():
    buf = io.BytesIO()
    doc = Document()
    doc.add_paragraph("Intro line for grading.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "B1"
    table.cell(1, 0).text = "A2"
    table.cell(1, 1).text = "B2"
    doc.save(buf)
    blocks = _classify_and_extract("hw.docx", buf.getvalue())
    text = _blocks_to_text(blocks)
    assert "Intro line" in text
    assert "A1" in text and "B2" in text
    assert "[表格]" in text


def test_extract_xlsx_sheet_rows():
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    assert ws is not None
    ws.title = "S1"
    ws.append(["Name", "Score"])
    ws.append(["Ann", 91])
    bio = io.BytesIO()
    wb.save(bio)
    blocks = _classify_and_extract("data.xlsx", bio.getvalue())
    text = _blocks_to_text(blocks)
    assert "工作表: S1" in text
    assert "Ann" in text and "91" in text


def test_extract_xls_fixture():
    from pathlib import Path

    p = Path(__file__).resolve().parents[2] / "fixtures" / "tiny.xls"
    assert p.is_file(), "fixtures/tiny.xls required (generated via xlwt)"
    blocks = _classify_and_extract("legacy.xls", p.read_bytes())
    text = _blocks_to_text(blocks)
    assert "HelloXls" in text
    assert "42" in text


def test_legacy_doc_non_ole_yields_empty():
    blocks = _classify_and_extract("fake.doc", b"not an ole file header")
    assert blocks == []


def test_zip_nested_xlsx():
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["Z", 1])
    inner = io.BytesIO()
    wb.save(inner)
    zbuf = io.BytesIO()
    with zipfile.ZipFile(zbuf, "w") as zf:
        zf.writestr("inner/table.xlsx", inner.getvalue())
    state: dict = {"file_count": 0, "total_bytes": 0, "skipped": []}
    blocks = _walk_zip_bytes(zbuf.getvalue(), root_path="bundle.zip", depth=1, state=state)
    text = _blocks_to_text(blocks)
    assert "Z" in text


def _rar_cli_available() -> bool:
    return shutil.which("rar") is not None


def test_rar_unencrypted_extracts_inner_txt(tmp_path):
    if not _rar_cli_available():
        pytest.skip("non-free `rar` CLI not installed (needed to build test .rar)")

    inner_zip = tmp_path / "inner.zip"
    with zipfile.ZipFile(inner_zip, "w") as zf:
        zf.writestr("note.txt", b"RAR_INNER_UNIQUE_TEXT_12345\n")

    rar_path = tmp_path / "t.rar"
    subprocess.run(
        ["rar", "a", "-ep", str(rar_path), str(inner_zip)],
        check=True,
        capture_output=True,
        cwd=str(tmp_path),
    )

    content = rar_path.read_bytes()
    state: dict = {"file_count": 0, "total_bytes": 0, "skipped": []}
    blocks = _walk_rar_bytes(content, root_path="t.rar", depth=1, state=state)
    text = _blocks_to_text(blocks)
    if not text and state.get("skipped"):
        reasons = " ".join(s.get("reason", "") for s in state["skipped"])
        if "unrar" in reasons.lower():
            pytest.skip("RAR unpack needs unrar: " + reasons)
    assert "RAR_INNER_UNIQUE_TEXT_12345" in text


def test_rar_password_rejected(tmp_path):
    if not _rar_cli_available():
        pytest.skip("non-free `rar` CLI not installed")

    solo = tmp_path / "solo.txt"
    solo.write_text("x", encoding="utf-8")
    rar_path = tmp_path / "enc.rar"
    subprocess.run(
        ["rar", "a", "-pPASS", str(rar_path), str(solo)],
        check=True,
        capture_output=True,
        cwd=str(tmp_path),
    )

    content = rar_path.read_bytes()
    state: dict = {"file_count": 0, "total_bytes": 0, "skipped": []}
    blocks = _walk_rar_bytes(content, root_path="enc.rar", depth=1, state=state)
    assert not blocks
    assert any("加密" in (s.get("reason") or "") for s in state["skipped"])
